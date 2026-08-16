# -*- coding: utf-8 -*-
"""生成提交用 Word 文档（.docx）：项目说明/解决方案、Demo 演示说明、技术架构与工作流程、
AI 工具/模型/素材及开源组件说明。

运行：python build_docx.py  → 输出到 提交材料/
依赖：python-docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "提交材料"
OUT.mkdir(exist_ok=True)

ACCENT = RGBColor(0xF4, 0x57, 0x1E)
TEXT = RGBColor(0x2B, 0x21, 0x1A)
MUTED = RGBColor(0x8A, 0x7D, 0x6D)
HEADER_FILL = "F4571E"

INTRO = ("小熊电器AI客服，基于大模型RAG，融合产品知识库与多平台规则，实现全自主接待，"
         "覆盖京东、天猫、抖音、拼多多四渠道，具备话术生成与多轮接待能力，人工转异常兜底，"
         "破解人力成本高、响应慢等五大痛点。")


def set_font(run, name="微软雅黑", size=10.5, bold=False, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, hexfill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexfill)
    tcpr.append(shd)


def H(doc, text, size=15, color=ACCENT, before=10, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    set_font(p.add_run(text), size=size, bold=True, color=color)
    return p


def P(doc, text, size=10.5, bold=False, color=TEXT, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def title(doc, main, sub):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(p.add_run(main), size=20, bold=True, color=ACCENT)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    set_font(p2.add_run(sub), size=10, color=MUTED)


def table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        set_font(c.paragraphs[0].add_run(h), size=9.5, bold=True,
                 color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(c, HEADER_FILL)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            set_font(cells[i].paragraphs[0].add_run(val), size=9.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def new_doc():
    d = Document()
    for s in d.sections:
        s.top_margin = Pt(40)
        s.bottom_margin = Pt(40)
        s.left_margin = Pt(48)
        s.right_margin = Pt(48)
    # 默认样式字体
    style = d.styles["Normal"]
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(10.5)
    return d


# ---------- 01 项目说明与解决方案 ----------
def doc_solution():
    d = new_doc()
    title(d, "小熊电器 AI 智能客服 —— 项目说明与解决方案",
          "基于大模型 RAG · 融合产品知识库与多平台规则 · 可替代人工客服")

    H(d, "一、一句话项目介绍（100 字以内）")
    P(d, INTRO, size=12, bold=True)
    P(d, f"（字数：{len(INTRO)} 字，含标点、字母、数字）", size=9, color=MUTED)

    H(d, "二、核心定位")
    P(d, "AI 智能客服作为「全自主接待单元」，直接面向买家进行售前咨询应答、产品推荐、售后问题处理，"
          "覆盖京东 / 天猫 / 抖音 / 拼多多 4 大电商渠道；人工客服转为异常兜底与升级处理。")
    P(d, "两条核心能力：", bold=True)
    P(d, "① 话术生成 —— 基于 RAG 检索产品知识 + 平台规则，实时生成贴合买家问题的回复话术。")
    P(d, "② 智能接待 —— 多轮对话上下文理解，主动引导、推荐、促单，覆盖「推荐 → 讲解 → 挖需 → 核心问答解决后促单」全链路。")
    P(d, "对应破解五大痛点：① 人力成本高 ② 响应时效受人力限制 ③ 标准化难统一 ④ 无 AI 自主接待能力 ⑤ 多平台规则适配难。")

    H(d, "三、技术架构（五层）")
    table(d, ["层", "职责", "实现"], [
        ["数据层", "产品知识、平台规则、样例咨询", "data/*.json（结构化、可替换）"],
        ["检索层", "从知识库召回与问题相关的产品", "BM25 + 中文字符 bigram 分词"],
        ["生成层", "基于检索结果 + 规则生成话术", "DeepSeek（OpenAI 兼容接口）"],
        ["对话层", "多轮上下文、意图识别、岗位路由", "conversation.py（售前/售中/售后）"],
        ["应用层", "可演示的 Web 界面 + API", "Streamlit + FastAPI"],
    ])
    P(d, "数据流：买家咨询 → 意图识别（售前/售中/售后）→ BM25 检索 → DeepSeek 生成 → 回复（含岗位/意图/命中）。")

    H(d, "四、工作流程（RAG）")
    P(d, "1. 意图识别：关键词规则分类（售前推荐 / 产品对比 / 参数 / 优惠 / 物流 / 售后 / 使用）。")
    P(d, "2. 检索 Retrieval：BM25 对产品文档打分，召回 Top-K 最相关产品；平台规则按当前渠道注入。")
    P(d, "3. 增强 Augmented：将「产品卖点/规格/竞品/FAQ」+「平台优惠/售后/话术风格」拼进 Prompt。")
    P(d, "4. 生成 Generation：DeepSeek 按「卖点 FAB + 对比 + 促单」三要素生成贴合平台风格的客服话术。")
    P(d, "多平台适配：一个模型 + 平台规则注入，四平台话术风格、优惠、售后自动切换，无需训练四套模型。")

    H(d, "五、制作说明")
    table(d, ["交付物", "制作方式"], [
        ["可运行代码 src/", "Python 3.13 + FastAPI + Streamlit，RAG 引擎（校验→分流→检索→生成→回写）"],
        ["演示视频 演示视频.mp4", "Pillow 逐帧渲染聊天动画 + edge-tts 中文配音 + FFmpeg 合成 1080p"],
        ["方案 PPT 方案PPT.pptx", "python-pptx 脚本生成（12 页，可编辑）"],
        ["交互 Demo 交互Demo.html", "纯 HTML/JS 离线版，内置意图识别 + 岗位分流 + 四平台预置话术"],
        ["冒烟测试 smoke_test.py", "不耗 API，验证检索→分流→生成链路可用"],
    ])

    H(d, "六、交付物清单")
    P(d, "src/（代码）、data/（样例知识库）、docs/（PPT / 视频 / 交互 Demo / 技术路线 / AI 组件说明）、"
          "README.md、requirements.txt、build_pptx.py、build_pdf.py、build_docx.py、make_demo_video.py、smoke_test.py。")
    P(d, "⚠️ .env（含真实密钥）不随提交包分发，仅提供 .env.example。", color=RGBColor(0xD1, 0x4A, 0x3C))

    d.save(str(OUT / "01_项目说明与解决方案.docx"))
    print("已生成 -> 01_项目说明与解决方案.docx")


# ---------- 05 Demo 演示说明 ----------
def doc_demo_guide():
    d = new_doc()
    title(d, "小熊电器 AI 智能客服 —— Demo 演示说明",
          "三种演示方式与现场操作步骤")

    H(d, "一、演示方式总览")
    table(d, ["方式", "载体", "适合场景"], [
        ["A 零安装交互", "docs/交互Demo.html", "现场双击即玩，无需联网 / Python / 密钥"],
        ["B 演示视频", "docs/演示视频.mp4", "无法现场操作时播放（101 秒）"],
        ["C 真实大模型", "Streamlit（src/app.py）", "展示真实 RAG + DeepSeek 生成"],
    ])

    H(d, "二、方式 A：交互 Demo（推荐现场演示）")
    P(d, "1. 双击打开「交互Demo.html」，浏览器自动进入。")
    P(d, "2. 顶部四个平台页签切换：京东 / 天猫 / 抖音 / 拼多多。")
    P(d, "3. 输入买家问题（或点下方示例），回车发送。")
    P(d, "4. 观察：AI 自动分流岗位（售前/售中/售后）、识别意图、命中产品，并按平台口吻回复。")
    P(d, "示例输入（对应岗位分流效果）：", bold=True)
    table(d, ["买家输入", "自动分流"], [
        ["想给爸妈买个养生壶，有推荐吗？", "售前 · 推荐"],
        ["这款按摩器和SKG比有什么优势？", "售前 · 对比"],
        ["现在下单有优惠吗？", "售中 · 优惠"],
        ["养生壶不想要了能退吗？", "售后 · 售后"],
    ])

    H(d, "三、方式 B：演示视频")
    P(d, "直接播放「演示视频.mp4」，共 6 个场景：开场定位 → 五大痛点 → 能力①话术生成 → 能力②智能接待全链路 → 跨平台售后 → 人工兜底收尾。")

    H(d, "四、方式 C：真实大模型")
    P(d, "1. pip install -r requirements.txt")
    P(d, "2. 复制 .env.example 为 .env，填入 DEEPSEEK_API_KEY")
    P(d, "3. streamlit run src/app.py（浏览器自动打开 http://localhost:8501）")
    P(d, "4. 三个 Tab：A 话术生成 / B 智能接待（多轮）/ C 话术质量对比。")

    H(d, "五、3 分钟演示话术")
    P(d, "开场（定位）：这是小熊电器 AI 智能客服，基于大模型 RAG，融合产品知识库与多平台规则，"
          "目标是可替代人工客服，实现全自主接待。")
    P(d, "痛点：人工客服有人力成本高、响应慢、话术难统一、无自主接待、多平台规则难适配五大痛点。")
    P(d, "能力①话术生成：买家问对比，AI 先检索命中产品知识，再实时生成卖点+对比+促单三要素话术。")
    P(d, "能力②智能接待：AI 多轮对话，主动推荐、讲解、挖需、促单，走完整条接待链路。")
    P(d, "跨渠道：四平台话术风格、优惠、售后自动适配，避免记混出错。")
    P(d, "收尾（人工兜底）：AI 自主接待售前/推荐/售后，人工转为异常兜底，五大痛点逐一解决。")

    d.save(str(OUT / "05_Demo演示说明.docx"))
    print("已生成 -> 05_Demo演示说明.docx")


# ---------- 06 技术架构与工作流程 ----------
def doc_arch():
    d = new_doc()
    title(d, "小熊电器 AI 智能客服 —— 技术架构与工作流程",
          "RAG 检索增强生成 · 五层架构 · 多平台适配")

    H(d, "一、整体架构（五层）")
    table(d, ["层", "职责", "实现"], [
        ["数据层", "产品知识、平台规则、样例咨询", "data/*.json（结构化、可替换）"],
        ["检索层", "从知识库召回与问题相关的产品", "BM25 + 中文字符 bigram 分词"],
        ["生成层", "基于检索结果 + 规则生成话术", "DeepSeek（OpenAI 兼容接口）"],
        ["对话层", "多轮上下文、意图识别、岗位路由", "conversation.py（售前/售中/售后）"],
        ["应用层", "可演示的 Web 界面 + API", "Streamlit + FastAPI"],
    ])

    H(d, "二、核心流程（RAG）")
    P(d, "1. 意图识别：对买家咨询做关键词规则分类（售前推荐 / 产品对比 / 参数 / 优惠 / 物流 / 售后 / 使用）。")
    P(d, "2. 检索 Retrieval：BM25 对产品文档打分，召回 Top-K 最相关产品；平台规则按当前接待渠道直接注入。")
    P(d, "3. 增强 Augmented：把「产品卖点 / 规格 / 竞品对比 / FAQ」和「平台优惠 / 售后 / 话术风格」拼进 Prompt 上下文。")
    P(d, "4. 生成 Generation：DeepSeek 按三要素（卖点 FAB + 对比 + 促单）生成贴合平台风格的客服话术。")

    H(d, "三、为什么选 BM25 而非向量检索")
    P(d, "· 产品库规模小（几十~几百 SKU），BM25 对关键词/品类/品牌/竞品名的召回已足够精准；")
    P(d, "· 零外部依赖、启动快、结果可解释，适合 Demo 快速跑通；")
    P(d, "· 预留升级路径：数据量变大后可平滑切换为向量检索（BGE 中文 embedding + 向量库）或 BM25 + 向量混合召回。")

    H(d, "四、多平台适配策略")
    P(d, "不训练四套模型，而是用「一个模型 + 平台规则注入」：")
    P(d, "· 把京东/天猫/抖音/拼多多的话术风格、优惠活动、售后政策、物流时效结构化进 platform_rules.json；")
    P(d, "· 生成时按当前渠道注入对应规则，模型自动切换语气（抖音偏热情口语、京东偏专业可靠）；")
    P(d, "· 平台规则可独立更新，无需重新训练。")

    H(d, "五、多轮对话与漏斗管理")
    P(d, "· 会话历史维护在 Conversation 对象中，生成时取最近 6 轮作为上下文；")
    P(d, "· 意图识别给出「当前漏斗阶段 + 下一步动作」提示，驱动完成 推荐 → 讲解 → 挖需 → 促单 全链路；")
    P(d, "· 意图识别当前用关键词规则，后续可替换为模型意图分类。")

    H(d, "六、生产化升级路线")
    table(d, ["阶段", "内容"], [
        ["MVP（当前）", "BM25 + DeepSeek + 规则意图 + Streamlit，可演示"],
        ["v1.1 向量化", "BGE 中文 embedding + 向量库，混合召回"],
        ["v1.2 评测体系", "话术质量评分集，自动化回归"],
        ["v1.3 意图升级", "模型意图分类 + 槽位抽取"],
        ["v2.0 生产化", "接入真实工单、知识库自动更新、人工兜底转接、成本监控"],
    ])

    d.save(str(OUT / "06_技术架构与工作流程.docx"))
    print("已生成 -> 06_技术架构与工作流程.docx")


# ---------- 07 AI 工具/模型/素材及开源组件说明（Word 版） ----------
def doc_ai():
    d = new_doc()
    title(d, "AI 工具、模型、素材及开源组件说明",
          "写明工具/模型名称及版本、主要用途、外部数据或素材来源、开源组件及授权情况；没有的项目填「无」")

    H(d, "一、大模型 / AI 工具")
    table(d, ["名称", "版本", "主要用途", "来源"], [
        ["DeepSeek（deepseek-v4-pro）", "v4", "话术生成（主模型，RAG 生成层）", "DeepSeek 官方 API（https://api.deepseek.com）"],
        ["DeepSeek（deepseek-v4-flash）", "v4", "话术生成（低延迟备选）", "DeepSeek 官方 API"],
    ])

    H(d, "二、开源组件")
    table(d, ["名称", "版本", "主要用途", "授权"], [
        ["Python", "3.13.12", "运行语言", "PSF License"],
        ["Streamlit", "1.61.1", "可视化 Demo 前端（三个 Tab）", "Apache-2.0"],
        ["FastAPI", "0.141.1", "HTTP API（/webhook /demo）", "MIT"],
        ["uvicorn", "0.52.3", "ASGI 服务器", "BSD-3-Clause"],
        ["pydantic", "2.13.4", "消息模型校验", "MIT"],
        ["python-dotenv", "1.2.2", "读取 .env 密钥", "BSD-3-Clause"],
        ["python-pptx", "1.0.2", "生成方案 PPT", "MIT"],
        ["Pillow", "12.2.0", "演示视频帧渲染", "HPND（Pillow License）"],
        ["edge-tts", "7.2.8", "演示视频中文配音", "MIT"],
        ["FFmpeg", "9.0（gyan.dev full build）", "视频/音频编码封装", "LGPL/GPL（本机处理，未再分发二进制）"],
        ["BM25 检索", "自研实现", "产品知识召回", "无（公开算法，自研零依赖）"],
    ])

    H(d, "三、外部数据 / 素材来源")
    table(d, ["项目", "说明", "来源"], [
        ["产品知识库 products.json", "8 款样例产品卖点/参数/竞品/FAQ", "自建样例（模拟小熊电器公开商品信息），正式使用需替换为官方商品中心数据"],
        ["平台规则 platform_rules.json", "四平台优惠/售后/物流/话术风格", "自建（基于各平台公开规则的一般描述）"],
        ["样例 QA sample_qa.json", "10 条咨询 + 人工参考话术", "自建样例"],
        ["图片 / 音乐 / 视频素材", "无（画面由代码绘制、配音由 TTS 合成）", "无"],
        ["字体", "微软雅黑 / 微软雅黑 Bold（系统自带）", "Windows 系统字体，仅本机渲染，未嵌入分发"],
        ["模型训练 / 微调", "无（仅 API 调用，未训练、未微调）", "无"],
    ])

    H(d, "四、结论")
    P(d, "① 大模型采用 DeepSeek 官方 API 调用，未训练、未微调；")
    P(d, "② 全部代码组件均为开源且许可友好（MIT / Apache-2.0 / BSD / PSF / HPND / LGPL），无商用授权冲突；")
    P(d, "③ 外部素材（图片、音乐、视频）无，画面与配音均为代码生成；")
    P(d, "④ 知识库为自建样例数据，正式上线前需替换为小熊电器官方授权数据。")

    d.save(str(OUT / "07_AI工具模型素材及开源组件说明.docx"))
    print("已生成 -> 07_AI工具模型素材及开源组件说明.docx")


if __name__ == "__main__":
    doc_solution()
    doc_demo_guide()
    doc_arch()
    doc_ai()
